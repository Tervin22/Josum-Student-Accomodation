from __future__ import annotations

from dataclasses import dataclass, field
from datetime import date
from pathlib import Path
from typing import Iterable
from zipfile import ZipFile, ZIP_DEFLATED

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT = "Josum Student Accommodation Management System"
CLIENT = "Josum Student Accommodation Management"
DEVELOPER = "Naxius Digital Solutions"
VERSION = "1.0"
STATUS = "Client Sign-Off Draft"
PACK_DATE = date.today().isoformat()
OUT_DIR = Path("docs") / f"josum-client-signoff-pack-{PACK_DATE}"
DOCS_DIR = OUT_DIR / "documents"

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(91, 103, 112)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"


@dataclass
class Section:
    title: str
    paragraphs: list[str] = field(default_factory=list)
    bullets: list[str] = field(default_factory=list)
    table: list[list[str]] | None = None
    table_widths: list[float] | None = None


@dataclass
class DocumentSpec:
    filename: str
    title: str
    document_id: str
    purpose: str
    sections: list[Section]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_in: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D7DBE2")


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{m}"))
        if element is None:
            element = OxmlElement(f"w:{m}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_header_footer(doc: Document, title: str) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run(f"{PROJECT} | {title}")
    set_run_font(run, size=8.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"{DEVELOPER} to {CLIENT} | {STATUS} | Page ")
    set_run_font(run, size=8.5, color=MUTED)
    add_page_number(footer)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, size=8.5, color=MUTED)


def add_title_block(doc: Document, spec: DocumentSpec) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("CLIENT SIGN-OFF PACK")
    set_run_font(run, size=10, color=MUTED, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(spec.title)
    set_run_font(run, size=24, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(PROJECT)
    set_run_font(run, size=12.5, color=MUTED)

    rows = [
        ["Prepared for", CLIENT, "Prepared by", DEVELOPER],
        ["Document ID", spec.document_id, "Version", VERSION],
        ["Date", PACK_DATE, "Status", STATUS],
        ["Classification", "Client confidential", "Approval purpose", "Review and sign-off"],
    ]
    add_table(doc, rows, widths=[1.25, 2.25, 1.25, 1.75], header=False)

    add_callout(doc, spec.purpose)
    doc.add_page_break()


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    set_cell_margins(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=NAVY, bold=True)


def add_control_sections(doc: Document, spec: DocumentSpec) -> None:
    doc.add_heading("Document Control", level=1)
    add_table(
        doc,
        [
            ["Version", "Date", "Prepared / Updated By", "Description"],
            [VERSION, PACK_DATE, DEVELOPER, "Initial client sign-off draft prepared for Josum review."],
        ],
        widths=[0.9, 1.2, 2.2, 2.2],
        header=True,
    )
    add_table(
        doc,
        [
            ["Role", "Name", "Organisation", "Signature", "Date"],
            ["Client Sponsor", "", CLIENT, "", ""],
            ["Client Operations Representative", "", CLIENT, "", ""],
            ["Developer Representative", "", DEVELOPER, "", ""],
        ],
        widths=[1.55, 1.45, 1.55, 1.3, 0.65],
        header=True,
    )

    doc.add_heading("Review Instructions", level=1)
    add_bullets(
        doc,
        [
            "Review the document for accuracy against the delivered system.",
            "Record any exception, clarification, or requested change before signing.",
            "Approval confirms acceptance of the scope described in this document, subject to any signed change requests.",
            "Credentials, registration keys, and operational secrets must be exchanged through a secure channel and are not embedded in this sign-off pack.",
        ],
    )
    doc.add_page_break()


def add_table(doc: Document, rows: list[list[str]], widths: list[float] | None = None, header: bool = True) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    set_cell_margins(table)
    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths and col_idx < len(widths):
                set_cell_width(cell, widths[col_idx])
            if header and row_idx == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(0)
            run = cell.paragraphs[0].add_run(str(value))
            set_run_font(run, size=9.2, color=NAVY if header and row_idx == 0 else None, bold=header and row_idx == 0)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_run_font(run, size=10.5)


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_run_font(run, size=10.5)


def add_body(doc: Document, spec: DocumentSpec) -> None:
    for section in spec.sections:
        doc.add_heading(section.title, level=1)
        for paragraph in section.paragraphs:
            p = doc.add_paragraph()
            run = p.add_run(paragraph)
            set_run_font(run, size=10.8)
        if section.bullets:
            add_bullets(doc, section.bullets)
        if section.table:
            add_table(doc, section.table, widths=section.table_widths, header=True)


def add_signoff(doc: Document, spec: DocumentSpec) -> None:
    doc.add_page_break()
    doc.add_heading("Client Acceptance and Sign-Off", level=1)
    p = doc.add_paragraph()
    p.add_run(
        f"By signing below, the authorised representatives confirm that {spec.title} has been reviewed and is accepted as a client sign-off document for {PROJECT}."
    )
    add_table(
        doc,
        [
            ["Decision", "Select"],
            ["Approved as written", ""],
            ["Approved with listed exceptions", ""],
            ["Not approved - changes required", ""],
        ],
        widths=[4.8, 1.7],
        header=True,
    )
    add_table(
        doc,
        [
            ["Name", "Role", "Organisation", "Signature", "Date"],
            ["", "Client authorised signatory", CLIENT, "", ""],
            ["", "Developer authorised signatory", DEVELOPER, "", ""],
        ],
        widths=[1.25, 1.75, 1.5, 1.3, 0.7],
        header=True,
    )
    doc.add_heading("Exceptions or Conditions", level=2)
    add_table(
        doc,
        [["Item", "Description", "Owner", "Target Date"], ["1", "", "", ""], ["2", "", "", ""]],
        widths=[0.6, 3.4, 1.25, 1.25],
        header=True,
    )


def save_doc(spec: DocumentSpec) -> Path:
    doc = Document()
    style_document(doc)
    add_header_footer(doc, spec.title)
    add_title_block(doc, spec)
    add_control_sections(doc, spec)
    add_body(doc, spec)
    add_signoff(doc, spec)
    path = DOCS_DIR / spec.filename
    doc.save(path)
    return path


def req_rows(prefix: str, rows: list[tuple[str, str, str, str]]) -> list[list[str]]:
    return [["ID", "Requirement", "Priority", "Acceptance Criteria"]] + [
        [f"{prefix}-{idx:03d}", req, priority, criteria] for idx, (req, priority, criteria) in enumerate(rows, start=1)
    ]


def common_scope() -> Section:
    return Section(
        "System Scope",
        bullets=[
            "Student registration, profile completion, application submission, document upload, offer acceptance, and self-service dashboard.",
            "Administrator and manager dashboards for applications, residences, rooms, maintenance, storage, finance reporting, inspections, communications, settings, and audit logs.",
            "Technician workflow for maintenance handling and SLA visibility.",
            "Security workflow for visitor check-in/check-out, visitor pre-registration lookup, incident reporting, and resident lookup.",
            "Room handover, check-in/check-out, storage, reporting, export, notifications, and role-based access control.",
        ],
    )


def specs() -> list[DocumentSpec]:
    return [
        DocumentSpec(
            "01_Business_Requirements_Document.docx",
            "Business Requirements Document",
            "JOSUM-BRD-001",
            "Defines the business objectives, scope, stakeholders, success criteria, and acceptance basis for the delivered student accommodation management system.",
            [
                common_scope(),
                Section(
                    "Business Objectives",
                    bullets=[
                        "Digitise the end-to-end residence application and administration process.",
                        "Reduce manual handling of student documents, room allocations, maintenance requests, visitor records, and storage forms.",
                        "Provide role-specific dashboards for students, administrators, managers, technicians, and security personnel.",
                        "Improve accountability through workflow statuses, audit logs, exportable reports, and sign-off records.",
                        "Support operational readiness for client go-live with maintainable, secure, and documented processes.",
                    ],
                ),
                Section(
                    "Stakeholders and Responsibilities",
                    table=[
                        ["Stakeholder", "Responsibility", "Expected Benefit"],
                        ["Josum Management", "Approve requirements, review dashboards, manage operational policies.", "Central view of operations and reporting."],
                        ["Josum Administration", "Review applications, documents, rooms, finance, inspections, communications, and settings.", "Reduced manual administration and improved traceability."],
                        ["Students", "Submit applications, upload documents, request maintenance/storage, pre-register visitors.", "Transparent self-service process."],
                        ["Security Staff", "Record visitor activity and incidents.", "Controlled access and clear shift records."],
                        ["Technicians", "Progress maintenance work and resolution notes.", "Focused maintenance queue and SLA visibility."],
                        ["Naxius Digital Solutions", "Deliver, document, support, and hand over the system.", "Clear completion and acceptance record."],
                    ],
                    table_widths=[1.55, 2.95, 2.0],
                ),
                Section(
                    "Success Criteria",
                    bullets=[
                        "Approved users can sign in and are routed to the correct role dashboard.",
                        "Student applications can be submitted, reviewed, approved, accepted, rejected, waitlisted, cancelled, or marked moved out.",
                        "Required student documents can be uploaded and downloaded by authorised users.",
                        "Room availability, occupancy, and allocation are visible to authorised staff.",
                        "Maintenance requests show acknowledgement and resolution SLA status.",
                        "Security, storage, inspections, finance, and communication workflows are usable from their dashboards.",
                        "CSV exports and audit logs support operational reporting.",
                    ],
                ),
                Section(
                    "Assumptions and Constraints",
                    bullets=[
                        "Client will provide accurate residence, room, rate, document, and operational policy data.",
                        "Client will appoint authorised staff for administrator, manager, technician, and security access.",
                        "Database backup, hosting, email provider, and operational support arrangements must be confirmed before production go-live.",
                        "Registration keys and passwords are controlled secrets and are excluded from sign-off documents.",
                    ],
                ),
            ],
        ),
        DocumentSpec(
            "02_Functional_Requirements_Specification.docx",
            "Functional Requirements Specification",
            "JOSUM-FRS-002",
            "Specifies the delivered functional behaviour, system modules, data handling, and acceptance criteria for each operational area.",
            [
                Section(
                    "Functional Requirement Matrix",
                    table=req_rows(
                        "FR",
                        [
                            ("The system shall authenticate users and route each user to the dashboard matching their role.", "Must", "Login succeeds for active users and redirects to Student, Admin, Manager, Security, or Technician dashboard."),
                            ("The system shall support staff registration through controlled role-specific registration authorisation.", "Must", "Only authorised staff registrations are accepted; secrets are not displayed in normal UI."),
                            ("Students shall submit applications with personal, academic, guarantor, next-of-kin, funding, medical, and declaration data.", "Must", "Application records are created with status SUBMITTED and appear in the student and admin dashboards."),
                            ("Administrators shall review, update, approve, reject, waitlist, cancel, and mark student applications moved out.", "Must", "Status changes are persisted, audited, and visible in history."),
                            ("Students shall upload required documents against an application.", "Must", "Files are stored securely, associated with the application, and downloadable by authorised users."),
                            ("Administrators shall manage residences, room types, and room availability.", "Must", "Operational dashboards show total, available, occupied, reserved, and maintenance room states."),
                            ("Students shall submit maintenance requests with category, priority, location, title, and description.", "Must", "Requests receive references, SLA deadlines, status OPEN, and notification records."),
                            ("Technicians and authorised staff shall update maintenance status and resolution notes.", "Must", "Acknowledgement/resolution timestamps and SLA state update correctly."),
                            ("Security shall record visitors, check-outs, incident reports, and student lookups.", "Must", "Visitor and incident records are searchable and auditable."),
                            ("Students shall submit storage requests using the storage-form fields and upload at least one item photograph.", "Must", "Storage request captures item count, site, student signature, and photos."),
                            ("Administrators/managers shall manage storage status and management signature.", "Must", "Storage status history and management signature are retained."),
                            ("Administrators/managers shall create inspection periods and room handover/check-in/check-out forms.", "Must", "Inspection records capture Botlogile form fields, documents, conditions, items, signatures, and exports."),
                            ("The system shall support communications and email template management.", "Should", "Authorised staff can send communications and manage templates."),
                            ("The system shall provide finance, storage, inspection, and other operational CSV exports.", "Must", "Exports download with current filters and audit records."),
                            ("The system shall record audit logs for material administrative actions.", "Must", "Audit logs show actor, action, entity, metadata, and timestamp."),
                        ],
                    ),
                    table_widths=[0.75, 3.35, 0.8, 1.6],
                ),
                Section(
                    "Non-Functional Functional Support",
                    bullets=[
                        "Forms shall validate required fields before submission.",
                        "Uploaded files shall be validated for extension, MIME type, file signature, and maximum file size.",
                        "Dashboards shall remain responsive on desktop and mobile screen sizes.",
                        "List views shall support search, filtering, pagination, and export where applicable.",
                    ],
                ),
            ],
        ),
        DocumentSpec(
            "03_User_Roles_and_Permissions_Matrix.docx",
            "User Roles and Permissions Matrix",
            "JOSUM-RBAC-003",
            "Defines system roles, access boundaries, and permission expectations for sign-off and operational governance.",
            [
                Section(
                    "Role Definitions",
                    table=[
                        ["Role", "Primary User", "Dashboard", "Access Intent"],
                        ["Student", "Resident/applicant", "Student dashboard", "Self-service applications, documents, maintenance, visitors, storage, profile."],
                        ["Administrator", "Josum admin team", "Administration dashboard", "Full operational management across modules, settings, templates, exports, audit."],
                        ["Manager", "Josum management", "Manager dashboard", "Operational oversight, reporting, storage, inspections, maintenance, finance visibility."],
                        ["Security", "Security personnel", "Security dashboard", "Visitor check-in/out, incident reporting, resident lookup."],
                        ["Technician", "Maintenance personnel", "Technician dashboard", "Maintenance queue, status updates, resolution notes, SLA visibility."],
                    ],
                    table_widths=[1.0, 1.45, 1.55, 2.5],
                ),
                Section(
                    "Permission Matrix",
                    table=[
                        ["Capability", "Student", "Administrator", "Manager", "Security", "Technician"],
                        ["Submit application", "Create/view own", "View/manage", "View reports", "No", "No"],
                        ["Upload application documents", "Create/view own", "Download/review", "View as authorised", "No", "No"],
                        ["Manage residences/rooms", "View public availability", "Full manage", "View/report", "No", "No"],
                        ["Maintenance request", "Create/view own", "Full manage", "Oversight", "No", "Assigned update"],
                        ["Maintenance SLA", "View own status", "Manage/escalate", "Monitor", "No", "View/update assigned"],
                        ["Visitor pre-registration", "Create/view own", "View", "View", "Process at gate", "No"],
                        ["Visitor check-in/check-out", "No", "View/admin", "View/report", "Create/update", "No"],
                        ["Incident report", "No", "View/manage", "View/report", "Create/update", "No"],
                        ["Storage request", "Create/view own", "Manage/export", "Manage/report", "No", "No"],
                        ["Room handover inspection", "No", "Create/manage/export", "Create/manage/report", "No", "No"],
                        ["Finance reporting", "No", "Export/view", "Export/view", "No", "No"],
                        ["Settings/email templates", "No", "Manage", "Limited/no direct edit unless granted", "No", "No"],
                        ["Audit logs", "No", "View", "View if authorised", "No", "No"],
                    ],
                    table_widths=[1.8, 1.05, 1.25, 1.05, 1.0, 1.35],
                ),
                Section(
                    "Access Control Rules",
                    bullets=[
                        "Users must be active to access protected system features.",
                        "Students can access only their own personal applications, documents, maintenance requests, storage requests, visitor pre-registrations, and profile records.",
                        "Staff registration is protected by role-specific authorisation and should be controlled by Josum management.",
                        "Role assignments and staff onboarding/offboarding must be approved by Josum management.",
                    ],
                ),
            ],
        ),
        DocumentSpec(
            "04_Application_Workflow_and_Status_Document.docx",
            "Application Workflow and Status Document",
            "JOSUM-WF-004",
            "Documents the lifecycle statuses and transitions used by the system across student applications and related operational workflows.",
            [
                Section(
                    "Application Status Workflow",
                    table=[
                        ["Status", "Meaning", "Typical Owner", "Allowed Next Steps"],
                        ["SUBMITTED", "Student has submitted an application.", "Student/Admin", "UNDER_REVIEW, APPROVED, REJECTED, WAITLISTED, CANCELLED"],
                        ["UNDER_REVIEW", "Administration is assessing the application and documents.", "Administrator", "APPROVED, REJECTED, WAITLISTED"],
                        ["APPROVED", "Application is approved and may be accepted by the student.", "Administrator/Student", "Accepted state through acceptedAt, CANCELLED, MOVED_OUT"],
                        ["REJECTED", "Application is declined.", "Administrator", "No normal forward transition; new application may be submitted if policy permits."],
                        ["WAITLISTED", "Application is held pending availability or review.", "Administrator", "APPROVED, REJECTED, CANCELLED"],
                        ["CANCELLED", "Student or administration has cancelled the application.", "Student/Admin", "No normal forward transition."],
                        ["MOVED_OUT", "Student has moved out and active accommodation is ended.", "Administrator", "Historical/reporting only."],
                    ],
                    table_widths=[1.2, 2.2, 1.3, 1.8],
                ),
                Section(
                    "Related Workflow Statuses",
                    table=[
                        ["Workflow", "Statuses"],
                        ["Maintenance", "OPEN, ACKNOWLEDGED, IN_PROGRESS, RESOLVED, CLOSED"],
                        ["Maintenance SLA", "ACK_PENDING, ACK_BREACHED, RESOLUTION_PENDING, RESOLUTION_BREACHED, RESOLVED"],
                        ["Storage", "DRAFT, SUBMITTED, UNDER_REVIEW, APPROVED, REJECTED, ITEMS_RECEIVED, ITEMS_RELEASED, CANCELLED"],
                        ["Inspections", "DRAFT, COMPLETED, FOLLOW_UP_REQUIRED, CLOSED"],
                        ["Room", "AVAILABLE, RESERVED, OCCUPIED, MAINTENANCE"],
                        ["Users", "ACTIVE, SUSPENDED"],
                    ],
                    table_widths=[1.5, 5.0],
                ),
                Section(
                    "Workflow Controls",
                    bullets=[
                        "Status changes should be performed only by authorised roles.",
                        "Material changes are audited with actor, status transition, entity, and timestamp.",
                        "Resolution or completion workflows require supporting notes or confirmation where configured.",
                        "Users should not bypass required document, room assignment, or acceptance steps unless approved by management.",
                    ],
                ),
            ],
        ),
        DocumentSpec(
            "05_Required_Documents_Matrix.docx",
            "Required Documents Matrix",
            "JOSUM-DOC-005",
            "Defines required student documents, upload expectations, review responsibility, and acceptance rules.",
            [
                Section(
                    "Required Document Matrix",
                    table=[
                        ["Document", "System Type", "Submitted By", "Required For", "Review Criteria"],
                        ["Applicant ID / Passport copy", "APPLICANT_ID_PASSPORT / ID_DOCUMENT", "Student", "Identity validation", "Readable, current, matches applicant identity."],
                        ["2 x Student color ID photos", "STUDENT_COLOR_ID_PHOTOS", "Student", "Residence file", "Clear color images suitable for resident profile/records."],
                        ["Student acceptance letter", "STUDENT_ACCEPTANCE_LETTER", "Student", "Academic eligibility", "Institution acceptance proof is readable and linked to applicant."],
                        ["Proof of registration / funding confirmation", "PROOF_OF_REGISTRATION", "Student", "Registration/funding verification", "Current academic year or accepted funding indication."],
                        ["Proof of payment", "PROOF_OF_PAYMENT", "Student", "Financial verification", "Payment proof is readable and date/reference is visible."],
                        ["Guarantor supporting documents", "GUARANTOR_SUPPORTING_DOCUMENTS", "Student/Guarantor", "Lease and payment assurance", "ID/passport copy, payslip, bank statement, and address proof where applicable."],
                        ["Medical aid certificate", "MEDICAL_AID_CERTIFICATE", "International students if applicable", "Medical coverage confirmation", "Valid certificate for required period."],
                        ["Other supporting document", "OTHER", "Student/Admin", "Exception handling", "Used only when requested by administration."],
                    ],
                    table_widths=[1.55, 1.45, 1.0, 1.25, 1.25],
                ),
                Section(
                    "Document Handling Rules",
                    bullets=[
                        "Document upload is linked to the relevant application and uploading user.",
                        "Authorised staff may download and review submitted documents.",
                        "Document completeness is used for operational review and reporting; missing documents should be communicated to the student.",
                        "Sensitive documents must be handled as confidential client/student data.",
                    ],
                ),
            ],
        ),
        DocumentSpec(
            "06_Data_Fields_and_Forms_Specification.docx",
            "Data Fields and Forms Specification",
            "JOSUM-DATA-006",
            "Documents the major data fields collected by the system forms for validation, support, and client sign-off.",
            [
                Section(
                    "Major Forms and Data Fields",
                    table=[
                        ["Form / Area", "Key Fields"],
                        ["Student profile", "First name, last name, email, phone, student number, institution, course, year of study, ID number, address, emergency contact, profile photo."],
                        ["Accommodation application", "Residence, room type, study year/semester, applicant category, registration status, student identity/contact, occupation date, nationality, gender, guarantor, next of kin, medical, funding, terms, declaration, electronic signature."],
                        ["Document upload", "Application, document type, original filename, storage key, MIME type, file size, checksum, uploader, timestamp."],
                        ["Maintenance request", "Category, priority, room type, location, title, description, status, acknowledgement/resolution deadlines, resolution note, assigned technician."],
                        ["Visitor pre-registration", "Visitor name, phone, ID number, relationship, expected date/time, host/residence/room, status."],
                        ["Visitor log", "Visitor details, host student, residence, room, reason, checked-in/out timestamps, security user, override reason where applicable."],
                        ["Incident report", "Incident type, severity, residence, room, description, status, reported/resolved metadata."],
                        ["Storage form", "Student name, student number, student signature, room number, item count, item descriptions, Josum One/Two site, storage notice acknowledgement, management signature, item photos."],
                        ["Room handover/check-in/check-out", "Student details, key number, check-in/out dates, document checklist, room condition move-in/out, items brought in, declarations and signatures."],
                    ],
                    table_widths=[1.85, 4.65],
                ),
                Section(
                    "Validation and Data Quality Rules",
                    bullets=[
                        "Required fields must be completed before submission.",
                        "Number fields must be converted and validated as numbers.",
                        "Boolean acknowledgements must be explicitly accepted where required.",
                        "File uploads are validated against allowed types, file signatures, and maximum upload size.",
                        "Searchable operational fields include references, student names, student numbers, rooms, statuses, and notes.",
                    ],
                ),
            ],
        ),
        DocumentSpec(
            "07_Room_Allocation_Rules_Document.docx",
            "Room Allocation Rules Document",
            "JOSUM-ROOM-007",
            "Documents the room availability, reservation, occupancy, and allocation rules used by administrators and management.",
            [
                Section(
                    "Room Allocation Rules",
                    table=[
                        ["Rule", "Description", "Control"],
                        ["Eligibility", "A student must have an application under review or approved before allocation.", "Administrator review."],
                        ["Availability", "Rooms should be allocated only when the selected room is available or otherwise authorised by management.", "Room status and dashboard availability."],
                        ["Residence and room type", "Allocation should match selected residence and room type preference where possible.", "Application and room records."],
                        ["Gender allocation", "Rooms may be filtered or allocated based on gender allocation rules configured for rooms.", "Room metadata."],
                        ["Acceptance", "Approved students accept an offer before storage and visitor resident workflows become available.", "acceptedAt timestamp."],
                        ["Occupancy", "Accepted room assignment contributes to occupancy and available-room reporting.", "Dashboard statistics."],
                        ["Move-out", "Moved-out applications should no longer count as active resident occupancy.", "Application status MOVED_OUT."],
                    ],
                    table_widths=[1.25, 3.55, 1.7],
                ),
                Section(
                    "Administrative Considerations",
                    bullets=[
                        "Room number, room name, residence, capacity, and status must be maintained accurately.",
                        "Manual overrides should be recorded in admin notes or audit trail where operationally significant.",
                        "Finance reporting should include approved, accepted, actively room-assigned residents only.",
                    ],
                ),
            ],
        ),
        DocumentSpec(
            "08_Maintenance_SLA_Agreement.docx",
            "Maintenance SLA Agreement",
            "JOSUM-SLA-008",
            "Defines the maintenance acknowledgement and resolution SLA rules implemented in the system.",
            [
                Section(
                    "SLA Targets",
                    table=[
                        ["SLA Area", "Target", "System Behaviour"],
                        ["Acknowledgement", "24 hours from request creation by default.", "Open requests show ACK_PENDING until acknowledged or breached."],
                        ["High/Urgent resolution", "12 hours from acknowledgement by default.", "High and urgent priorities use the high-resolution SLA value."],
                        ["Low/Medium resolution", "48 hours from acknowledgement by default.", "Low and medium priorities use the low-resolution SLA value."],
                        ["SLA sweep", "Every 5 minutes by default.", "Background process updates breached SLA statuses and sends reminders."],
                    ],
                    table_widths=[1.5, 2.0, 3.0],
                ),
                Section(
                    "Maintenance Status Workflow",
                    table=[
                        ["Status", "SLA Meaning", "Operational Action"],
                        ["OPEN", "Awaiting acknowledgement.", "Administrator/technician acknowledges and starts work."],
                        ["ACKNOWLEDGED", "Resolution timer active.", "Technician investigates and updates status."],
                        ["IN_PROGRESS", "Resolution timer active.", "Technician continues work."],
                        ["RESOLVED", "SLA complete.", "Resolution note required."],
                        ["CLOSED", "SLA complete and administratively closed.", "No further action unless reopened outside system scope."],
                    ],
                    table_widths=[1.25, 2.25, 3.0],
                ),
                Section(
                    "SLA Governance",
                    bullets=[
                        "SLA targets may be changed through environment configuration before production deployment.",
                        "Resolution or closure requires a resolution note.",
                        "SLA breaches trigger system/email notifications to responsible operational users.",
                        "SLA reporting should be reviewed by management during operational meetings.",
                    ],
                ),
            ],
        ),
        DocumentSpec(
            "09_Notification_and_Email_Matrix.docx",
            "Notification and Email Matrix",
            "JOSUM-NOTIFY-009",
            "Defines expected system and email notifications for operational events and client acceptance.",
            [
                Section(
                    "Notification Matrix",
                    table=[
                        ["Event", "Recipient", "Channel", "Purpose"],
                        ["Student registration/application submission", "Student / Administration", "System and email where configured", "Confirm submission and alert review team."],
                        ["Application status change", "Student", "System/email", "Inform student of approval, rejection, waitlist, cancellation, or move-out state."],
                        ["Document upload", "Administration", "System/email where configured", "Alert staff that supporting evidence is available."],
                        ["Maintenance submitted", "Student and maintenance/admin users", "System/email", "Confirm receipt and create operational queue item."],
                        ["Maintenance status changed", "Student", "System/email", "Communicate acknowledgement, progress, resolution, or closure."],
                        ["Maintenance SLA breach", "Responsible operations users", "System/email", "Escalate acknowledgement or resolution breach."],
                        ["Storage request submitted", "Student", "System/email", "Confirm storage request reference and review start."],
                        ["Storage status changed", "Student", "System/email", "Notify approval, rejection, receipt, release, or cancellation."],
                        ["Visitor pre-registration", "Student/security", "System", "Support expected visitor verification."],
                        ["Communications broadcast", "Selected recipients", "Email/system", "Send operational notices from authorised staff."],
                    ],
                    table_widths=[1.75, 1.5, 1.35, 1.9],
                ),
                Section(
                    "Template Governance",
                    bullets=[
                        "Email templates should be approved by Josum before go-live.",
                        "Templates should avoid exposing sensitive operational secrets.",
                        "Failed email sends should not block core record creation but should be logged for support follow-up.",
                    ],
                ),
            ],
        ),
        DocumentSpec(
            "10_Reporting_and_Export_Specification.docx",
            "Reporting and Export Specification",
            "JOSUM-REPORT-010",
            "Defines dashboards, reports, filters, and export expectations for operational and management reporting.",
            [
                Section(
                    "Reports and Exports",
                    table=[
                        ["Report / Export", "Users", "Filters / Scope", "Output"],
                        ["Application dashboard", "Admin/Manager", "Status, search, residence, rooms", "On-screen table and operational views."],
                        ["Finance export", "Admin/Manager", "Approved, accepted, active residents", "CSV with identity, funding, contact, accommodation, and acceptance details."],
                        ["Storage export", "Admin/Manager", "Search, status, residence", "CSV with student, storage site, item count, signatures, dates, notes, file count."],
                        ["Inspection export", "Admin/Manager", "Search, status, period, residence", "CSV with handover fields, condition fields, signatures, photos, follow-up."],
                        ["Maintenance dashboard", "Admin/Manager/Technician", "Status, priority, category, search", "On-screen queue with SLA state and action controls."],
                        ["Audit logs", "Admin/Manager if authorised", "Action/entity/date search", "On-screen compliance trail."],
                    ],
                    table_widths=[1.45, 1.35, 2.1, 1.6],
                ),
                Section(
                    "Export Controls",
                    bullets=[
                        "Exports must be restricted to authorised roles.",
                        "Exports should use current filters where available.",
                        "Export activity should be auditable because exported files may contain personal information.",
                        "CSV output is intended for Excel-compatible operational analysis.",
                    ],
                ),
            ],
        ),
        DocumentSpec(
            "11_Security_and_Privacy_Requirements_Document.docx",
            "Security and Privacy Requirements Document",
            "JOSUM-SEC-011",
            "Defines security, privacy, access control, file handling, and operational protection requirements for sign-off.",
            [
                Section(
                    "Security Requirements",
                    table=[
                        ["Area", "Requirement"],
                        ["Authentication", "Protected routes require valid authenticated sessions and token handling."],
                        ["Authorisation", "Role-based access controls restrict each user to approved functions."],
                        ["Passwords", "Passwords must be stored as secure hashes, not plain text."],
                        ["Staff registration", "Role registration must be controlled by authorised keys or administrative approval; keys must not be exposed in documents."],
                        ["File uploads", "Files are validated by extension, MIME type, content signature, file size, and allowed type."],
                        ["Audit", "Material administrative actions are captured in audit logs."],
                        ["Data minimisation", "Only necessary personal and operational data should be collected for accommodation management."],
                        ["Confidentiality", "Student IDs, contact details, funding, guarantor documents, and medical details must be treated as confidential."],
                    ],
                    table_widths=[1.55, 4.95],
                ),
                Section(
                    "Privacy and Operational Controls",
                    bullets=[
                        "Client should define retention periods for applications, documents, visitor logs, incident reports, and maintenance records.",
                        "Client should nominate system owners responsible for access review and staff offboarding.",
                        "Backups must be encrypted or protected according to the hosting environment selected for production.",
                        "Production credentials must be stored outside source control and shared through a secure channel.",
                    ],
                ),
            ],
        ),
        DocumentSpec(
            "12_UI_UX_Design_Approval.docx",
            "UI/UX Design Approval",
            "JOSUM-UIUX-012",
            "Provides a formal approval template for the delivered user experience, dashboards, navigation, and role-specific screens.",
            [
                Section(
                    "UI/UX Review Scope",
                    table=[
                        ["Screen / Experience", "Review Focus", "Approval Criteria"],
                        ["Public pages", "Residence browsing and navigation.", "Clear branding, readable content, responsive layout."],
                        ["Login/register", "Role registration, staff/student flows.", "Correct validation and role routing."],
                        ["Student dashboard", "Applications, documents, maintenance, visitors, storage, profile.", "Workflows are understandable and forms are complete."],
                        ["Admin dashboard", "Operational tabs and dense tables.", "Staff can search, filter, update, export, and understand statuses."],
                        ["Manager dashboard", "Management reporting and oversight.", "Reports and summaries support decision-making."],
                        ["Security dashboard", "Visitor and incident operation.", "Fast entry, clear status, and resident lookup."],
                        ["Technician dashboard", "Maintenance queue and updates.", "Technicians can progress assigned work efficiently."],
                    ],
                    table_widths=[1.5, 2.4, 2.6],
                ),
                Section(
                    "Approval Checklist",
                    bullets=[
                        "Branding is acceptable for Josum Student Accommodation.",
                        "Navigation labels are clear to target users.",
                        "Forms are understandable, complete, and not visually overloaded.",
                        "Mobile and desktop layouts are acceptable for operational use.",
                        "Error and success messages are clear enough for users.",
                    ],
                ),
            ],
        ),
        DocumentSpec(
            "13_Data_Migration_Approval.docx",
            "Data Migration Approval",
            "JOSUM-MIG-013",
            "Provides the template and approval record for any migration of Josum data into the production system.",
            [
                Section(
                    "Migration Scope",
                    table=[
                        ["Data Area", "Source", "Target", "Migration Required"],
                        ["Users/students", "Client-provided records or manual registration", "User and StudentProfile", "To be confirmed by client."],
                        ["Residences/rooms", "Client room list", "Residence, ResidenceRoom, RoomType", "Yes if bulk data exists."],
                        ["Applications", "Legacy application files if applicable", "Application", "Optional; otherwise new submissions only."],
                        ["Documents", "Legacy file folders if applicable", "Document storage", "Optional and subject to data quality."],
                        ["Maintenance/history", "Legacy logs if applicable", "MaintenanceRequest", "Optional."],
                    ],
                    table_widths=[1.4, 1.75, 1.75, 1.6],
                ),
                Section(
                    "Migration Approval Criteria",
                    bullets=[
                        "Client approves source data files and confirms they are complete.",
                        "Mapping from source columns to system fields is reviewed and signed.",
                        "Sample migrated records are validated by Josum before final import.",
                        "Backup and rollback approach is confirmed before production import.",
                        "Migration completion is documented with record counts and known exceptions.",
                    ],
                ),
                Section(
                    "Migration Sign-Off Record",
                    table=[
                        ["Checkpoint", "Accepted", "Comments"],
                        ["Source data received", "", ""],
                        ["Field mapping approved", "", ""],
                        ["Test import validated", "", ""],
                        ["Production import approved", "", ""],
                        ["Post-import reconciliation accepted", "", ""],
                    ],
                    table_widths=[2.4, 1.0, 3.1],
                ),
            ],
        ),
        DocumentSpec(
            "14_User_Acceptance_Testing_Plan.docx",
            "User Acceptance Testing Plan",
            "JOSUM-UAT-014",
            "Defines the client-led UAT approach, test coverage, entry criteria, exit criteria, roles, and test cases.",
            [
                Section(
                    "UAT Approach",
                    bullets=[
                        "UAT is conducted by Josum representatives using realistic sample users and records.",
                        "Testing covers each user role and each critical business workflow.",
                        "Defects are logged with severity, screenshots, expected result, actual result, and owner.",
                        "UAT exit requires all critical/high defects resolved or formally deferred by the client.",
                    ],
                ),
                Section(
                    "Core UAT Test Cases",
                    table=[
                        ["ID", "Scenario", "Expected Result"],
                        ["UAT-01", "Student registers, logs in, updates profile, and uploads profile photo.", "Student dashboard is accessible and profile is saved."],
                        ["UAT-02", "Student submits an accommodation application with declaration/signature.", "Application appears in admin review queue."],
                        ["UAT-03", "Student uploads all required documents.", "Documents are listed and downloadable by admin."],
                        ["UAT-04", "Admin approves application and assigns/validates room.", "Student sees approved application and can accept."],
                        ["UAT-05", "Student logs maintenance request; technician progresses to resolved.", "SLA and status update correctly with resolution note."],
                        ["UAT-06", "Security checks in and checks out a visitor.", "Visitor log shows correct timestamps and security user."],
                        ["UAT-07", "Student submits storage form with item photos.", "Storage request is visible to admin/manager with files."],
                        ["UAT-08", "Admin creates check-in/check-out handover inspection.", "Inspection stores form fields and exports correctly."],
                        ["UAT-09", "Manager downloads finance, storage, and inspection exports.", "CSV files open with expected fields."],
                        ["UAT-10", "Admin reviews audit logs and settings/email templates.", "Authorised data is visible and update actions are controlled."],
                    ],
                    table_widths=[0.85, 3.6, 2.05],
                ),
                Section(
                    "Entry and Exit Criteria",
                    table=[
                        ["Criteria Type", "Criteria"],
                        ["Entry", "System deployed to UAT environment, test users available, sample data loaded, client testers identified."],
                        ["Exit", "Core UAT cases passed, critical/high defects closed or deferred, sign-off completed by authorised client representative."],
                    ],
                    table_widths=[1.2, 5.3],
                ),
            ],
        ),
        DocumentSpec(
            "15_User_Acceptance_Testing_Sign-Off.docx",
            "User Acceptance Testing Sign-Off",
            "JOSUM-UAT-SO-015",
            "Provides the formal record for Josum to accept UAT results and authorise progression toward go-live.",
            [
                Section(
                    "UAT Result Summary",
                    table=[
                        ["Metric", "Result / Notes"],
                        ["UAT period", ""],
                        ["Number of test cases planned", ""],
                        ["Number of test cases passed", ""],
                        ["Open critical defects", ""],
                        ["Open high defects", ""],
                        ["Deferred items/change requests", ""],
                        ["Overall UAT decision", ""],
                    ],
                    table_widths=[2.2, 4.3],
                ),
                Section(
                    "Sign-Off Conditions",
                    bullets=[
                        "All mandatory UAT scenarios have either passed or been accepted with documented exceptions.",
                        "The client understands any deferred items will require change request approval or post-go-live support scheduling.",
                        "This sign-off authorises the project team to proceed to go-live readiness activities.",
                    ],
                ),
                Section(
                    "Defect Exceptions",
                    table=[
                        ["Defect / Item", "Severity", "Decision", "Owner", "Target Date"],
                        ["", "", "", "", ""],
                        ["", "", "", "", ""],
                    ],
                    table_widths=[2.0, 0.9, 1.35, 1.15, 1.1],
                ),
            ],
        ),
        DocumentSpec(
            "16_Change_Request_Form.docx",
            "Change Request Form",
            "JOSUM-CR-016",
            "Provides the formal template for requesting, assessing, approving, and tracking changes after baseline approval.",
            [
                Section(
                    "Change Request Details",
                    table=[
                        ["Field", "Response"],
                        ["Change request ID", ""],
                        ["Requested by", ""],
                        ["Date requested", ""],
                        ["Module / area affected", ""],
                        ["Current approved behaviour", ""],
                        ["Requested change", ""],
                        ["Business reason", ""],
                        ["Priority", ""],
                    ],
                    table_widths=[2.0, 4.5],
                ),
                Section(
                    "Impact Assessment",
                    table=[
                        ["Impact Area", "Assessment"],
                        ["Functional scope", ""],
                        ["Data/model impact", ""],
                        ["Security/privacy impact", ""],
                        ["UI/UX impact", ""],
                        ["Testing required", ""],
                        ["Estimated effort/cost", ""],
                        ["Target release", ""],
                    ],
                    table_widths=[2.0, 4.5],
                ),
                Section(
                    "Approval Decision",
                    table=[
                        ["Decision", "Select", "Comments"],
                        ["Approved", "", ""],
                        ["Rejected", "", ""],
                        ["Deferred", "", ""],
                    ],
                    table_widths=[1.5, 1.0, 4.0],
                ),
            ],
        ),
        DocumentSpec(
            "17_Go-Live_Approval.docx",
            "Go-Live Approval",
            "JOSUM-GO-017",
            "Provides the operational go-live readiness checklist and client approval record.",
            [
                Section(
                    "Go-Live Readiness Checklist",
                    table=[
                        ["Readiness Area", "Criteria", "Accepted"],
                        ["UAT", "UAT sign-off completed and critical defects closed/deferred.", ""],
                        ["Production environment", "Hosting, database, email, storage, and environment variables configured.", ""],
                        ["Data", "Production seed/migration/import completed and reconciled.", ""],
                        ["Security", "Admin/staff access confirmed; secrets exchanged securely; default/test access removed.", ""],
                        ["Backups", "Backup and restore responsibility confirmed.", ""],
                        ["Training", "Key users briefed on dashboards and workflows.", ""],
                        ["Support", "Support contacts, escalation path, and warranty/support window confirmed.", ""],
                        ["Communication", "Go-live date and operational change communicated to staff/users.", ""],
                    ],
                    table_widths=[1.55, 4.15, 0.8],
                ),
                Section(
                    "Go/No-Go Decision",
                    bullets=[
                        "Go-live approval authorises the system to be made available for production users.",
                        "Any approved exceptions must be listed with owner and target date.",
                        "Client is responsible for production business process adoption and staff access governance after go-live.",
                    ],
                ),
            ],
        ),
        DocumentSpec(
            "18_Project_Completion_and_Handover_Sign-Off.docx",
            "Project Completion and Handover Sign-Off",
            "JOSUM-HANDOVER-018",
            "Confirms completion of the project deliverables, handover items, and client acceptance responsibilities.",
            [
                Section(
                    "Handover Items",
                    table=[
                        ["Item", "Description", "Status"],
                        ["Application source/workspace", "System files and implementation workspace handed to client-designated location.", ""],
                        ["Database schema", "Prisma/PostgreSQL schema and migrations included.", ""],
                        ["Environment configuration", "Environment variable template and production values to be controlled securely.", ""],
                        ["Run scripts", "Local start script and build commands provided.", ""],
                        ["Documentation pack", "BRD, FRS, matrices, approvals, UAT, go-live, and handover documents.", ""],
                        ["Operational training", "Admin/manager/security/technician/student workflow orientation.", ""],
                        ["Support process", "Post-handover support contact and change request process confirmed.", ""],
                    ],
                    table_widths=[1.65, 3.8, 1.05],
                ),
                Section(
                    "Completion Criteria",
                    bullets=[
                        "All agreed modules are delivered for client review.",
                        "Client sign-off documents are provided.",
                        "Production deployment/go-live is approved or pending only listed exceptions.",
                        "Client accepts responsibility for operational data, authorised user management, and post-go-live business process execution.",
                    ],
                ),
                Section(
                    "Final Handover Exceptions",
                    table=[
                        ["Exception", "Owner", "Target Date", "Resolution Notes"],
                        ["", "", "", ""],
                        ["", "", "", ""],
                    ],
                    table_widths=[2.3, 1.2, 1.0, 2.0],
                ),
            ],
        ),
    ]


def create_index(doc_paths: list[Path]) -> Path:
    spec = DocumentSpec(
        "00_Client_Sign-Off_Pack_Index.docx",
        "Client Sign-Off Pack Index",
        "JOSUM-SIGNOFF-000",
        "Provides the document register for the complete client sign-off pack issued by Naxius Digital Solutions to Josum Student Accommodation Management.",
        [
            Section(
                "Document Register",
                table=[["No.", "Document", "Filename", "Purpose"]]
                + [
                    [f"{idx:02d}", path.stem.replace("_", " "), path.name, "Review and sign-off"]
                    for idx, path in enumerate(doc_paths, start=1)
                ],
                table_widths=[0.55, 2.1, 2.25, 1.6],
            ),
            Section(
                "Recommended Client Review Order",
                bullets=[
                    "Review core requirements documents first: BRD, FRS, roles, workflows, documents, data fields, and room rules.",
                    "Review operational controls next: maintenance SLA, notifications, reports, security/privacy, and UI/UX approval.",
                    "Review delivery controls last: data migration, UAT plan, UAT sign-off, change request, go-live, and handover sign-off.",
                    "Sign only after exceptions have been listed or accepted through the change request process.",
                ],
            ),
        ],
    )
    return save_doc(spec)


def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    generated = [save_doc(spec) for spec in specs()]
    index = create_index(generated)
    all_docs = [index] + generated
    zip_path = OUT_DIR / f"Josum_Client_Signoff_Pack_{PACK_DATE}.zip"
    with ZipFile(zip_path, "w", ZIP_DEFLATED) as archive:
        for path in all_docs:
            archive.write(path, arcname=path.name)
    print(zip_path)
    for path in all_docs:
        print(path)


if __name__ == "__main__":
    main()
