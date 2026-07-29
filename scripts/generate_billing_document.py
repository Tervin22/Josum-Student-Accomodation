from __future__ import annotations

from datetime import date, timedelta
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


TODAY = date.today()
DUE_DATE = TODAY + timedelta(days=7)
OUT_DIR = Path("docs") / "billing"
OUT_PATH = OUT_DIR / f"Naxius_to_Josum_Billing_Document_{TODAY.isoformat()}.docx"

DEVELOPER = "Naxius Digital Solutions"
CLIENT = "Josum Student Accommodation Management"
PROJECT = "Josum Student Accommodation Management System"
DOC_NO = f"NDS-JOSUM-BILL-{TODAY:%Y%m%d}-001"

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(35, 45, 58)
MUTED = RGBColor(91, 103, 112)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_in: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D7DBE2") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def add_table(doc, rows, widths=None, header=True):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    set_cell_margins(table)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths and c_idx < len(widths):
                set_cell_width(cell, widths[c_idx])
            if header and r_idx == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(value))
            set_run_font(run, size=9.2, color=NAVY if header and r_idx == 0 else DARK, bold=header and r_idx == 0)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
    set_run_font(run, size=8.5, color=MUTED)


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 15, BLUE, 12, 6),
        ("Heading 2", 12.5, BLUE, 8, 4),
        ("Heading 3", 11.5, DARK, 6, 3),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    run = header.add_run(f"{DEVELOPER} | Billing Document | {PROJECT}")
    set_run_font(run, size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"{DOC_NO} | Client Confidential | Page ")
    set_run_font(run, size=8.5, color=MUTED)
    add_page_number(footer)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(DEVELOPER)
    set_run_font(run, size=11.5, color=MUTED, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("BILLING DOCUMENT")
    set_run_font(run, size=25, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(f"Project billing statement for {PROJECT}")
    set_run_font(run, size=12.5, color=MUTED)

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="B7C5D6")
    set_cell_margins(table, top=110, bottom=110, start=140, end=140)
    left, right = table.rows[0].cells
    set_cell_width(left, 3.65)
    set_cell_width(right, 2.85)
    set_cell_shading(left, LIGHT_BLUE)
    set_cell_shading(right, WHITE)

    left.paragraphs[0].add_run("Billed To").bold = True
    left.add_paragraph(CLIENT)
    left.add_paragraph("Client contact: [Insert client representative]")
    left.add_paragraph("Client address: [Insert client billing address]")

    right.paragraphs[0].add_run("Invoice Details").bold = True
    right.add_paragraph(f"Document No.: {DOC_NO}")
    right.add_paragraph(f"Issue Date: {TODAY.isoformat()}")
    right.add_paragraph(f"Due Date: {DUE_DATE.isoformat()}")
    right.add_paragraph("Currency: South African Rand (ZAR)")


def add_notice(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="B7C5D6")
    set_cell_margins(table, top=100, bottom=100, start=140, end=140)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF8E8")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(
        "Note: This billing document is prepared as an editable client billing statement. "
        "Insert the agreed commercial values, VAT status, banking details, and company registration/tax details before final issue."
    )
    set_run_font(run, size=9.5, color=RGBColor(122, 90, 0), bold=True)


def add_body(doc: Document) -> None:
    doc.add_heading("Supplier and Client Information", level=1)
    add_table(
        doc,
        [
            ["Field", "Supplier", "Client"],
            ["Organisation", DEVELOPER, CLIENT],
            ["Project", PROJECT, PROJECT],
            ["Contact person", "[Insert Naxius representative]", "[Insert Josum representative]"],
            ["Email", "[Insert Naxius email]", "[Insert Josum email]"],
            ["Address", "[Insert Naxius business address]", "[Insert Josum billing address]"],
            ["Registration / VAT", "[Insert company reg. and VAT details if applicable]", "[Insert client billing/VAT details if applicable]"],
        ],
        widths=[1.4, 2.55, 2.55],
    )

    doc.add_heading("Billing Summary", level=1)
    add_table(
        doc,
        [
            ["Description", "Billing Basis", "Amount"],
            ["Student Accommodation Management System development and implementation", "Fixed project / agreed commercial scope", "R [Insert amount]"],
            ["Subtotal", "", "R [Insert subtotal]"],
            ["VAT, if applicable", "[Insert VAT rate or N/A]", "R [Insert VAT]"],
            ["Total Due", "", "R [Insert total due]"],
            ["Amount Paid / Deposit", "If applicable", "R [Insert paid amount]"],
            ["Balance Due", "", "R [Insert balance due]"],
        ],
        widths=[3.4, 1.8, 1.3],
    )

    doc.add_heading("Itemised Deliverables", level=1)
    add_table(
        doc,
        [
            ["No.", "Deliverable / Service", "Billing Status", "Amount"],
            ["1", "Requirements analysis, planning, and production-style system preparation.", "Included", "R [Insert]"],
            ["2", "Student application, document upload, profile, and student dashboard workflows.", "Included", "R [Insert]"],
            ["3", "Administrator and manager operational dashboards, reporting, filtering, and exports.", "Included", "R [Insert]"],
            ["4", "Residence, room, room allocation, occupancy, finance reporting, and room management features.", "Included", "R [Insert]"],
            ["5", "Maintenance module, technician dashboard, SLA tracking, reminders, and resolution workflow.", "Included", "R [Insert]"],
            ["6", "Security module including visitor check-in/check-out, resident lookup, visitor pre-registration, and incident reporting.", "Included", "R [Insert]"],
            ["7", "Student storage process using Josum storage-form fields and item photograph uploads.", "Included", "R [Insert]"],
            ["8", "Room handover, check-in/check-out, inspection periods, condition tracking, signatures, and inspection exports.", "Included", "R [Insert]"],
            ["9", "Notifications, email templates, audit logs, system clean-up, production readiness, and role routing.", "Included", "R [Insert]"],
            ["10", "Client sign-off documentation pack, UAT/go-live templates, and handover documentation.", "Included", "R [Insert]"],
        ],
        widths=[0.45, 3.95, 1.15, 0.95],
    )

    doc.add_heading("Payment Terms", level=1)
    add_table(
        doc,
        [
            ["Term", "Detail"],
            ["Payment due", f"Within 7 calendar days of issue date unless otherwise agreed in writing. Due date: {DUE_DATE.isoformat()}."],
            ["Payment reference", DOC_NO],
            ["Late payment", "Any late-payment interest, suspension, or support limitation terms should be inserted if contractually agreed."],
            ["Scope note", "This billing document covers delivered system development and handover work for the Josum Student Accommodation Management System."],
            ["Change requests", "Any additional work requested after sign-off should be billed through an approved change request or support agreement."],
        ],
        widths=[1.65, 4.85],
    )

    doc.add_heading("Banking Details", level=1)
    add_table(
        doc,
        [
            ["Field", "Detail"],
            ["Account holder", "[Insert Naxius Digital Solutions account holder]"],
            ["Bank", "[Insert bank name]"],
            ["Account number", "[Insert account number]"],
            ["Branch code", "[Insert branch code]"],
            ["Account type", "[Insert account type]"],
            ["Payment reference", DOC_NO],
        ],
        widths=[1.65, 4.85],
    )

    doc.add_heading("Billing Notes", level=1)
    notes = [
        "All amounts must be confirmed against the approved commercial agreement before this document is issued as a final invoice.",
        "This document excludes passwords, registration keys, database credentials, or other operational secrets.",
        "Client sign-off documents and handover documentation have been prepared separately for acceptance and project closure.",
        "Support, hosting, maintenance, future enhancements, and data migration beyond the agreed scope should be billed separately unless included in the commercial agreement.",
    ]
    for note in notes:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(note)
        set_run_font(run, size=10.2, color=DARK)

    doc.add_heading("Client Payment Acknowledgement", level=1)
    add_table(
        doc,
        [
            ["Name", "Role", "Organisation", "Signature", "Date"],
            ["", "Client authorised representative", CLIENT, "", ""],
            ["", "Naxius authorised representative", DEVELOPER, "", ""],
        ],
        widths=[1.25, 1.75, 1.5, 1.3, 0.7],
    )


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_document(doc)
    add_title(doc)
    add_notice(doc)
    add_body(doc)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
