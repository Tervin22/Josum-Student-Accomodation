from __future__ import annotations

import re
import shutil
import tempfile
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image
from lxml import etree


SOURCE_DOCX = Path(r"C:\Users\WalterTervinShingube\Downloads\Functional Requirements Specification (FRS).docx")
SOURCE_LOGO = Path(r"C:\Users\WalterTervinShingube\Downloads\ChatGPT Image Jul 28, 2026, 03_52_51 PM.png")
OUT_DIR = Path(r"C:\Users\WalterTervinShingube\Downloads\LGSETA\nathi-student-accommodation")
OUTPUT_DOCX = OUT_DIR / "Functional Requirements Specification (FRS) - Redesigned.docx"
LOGO_CROP = OUT_DIR / "synexis-logo-horizontal.png"


PRIMARY_BLUE = "0878E8"
DEEP_NAVY = "082888"
CHARCOAL = "283038"
DARK_CHARCOAL = "202832"
LIGHT_BLUE = "D8E8F8"
LIGHT_GRAY = "E8E8E8"
WHITE = "FFFFFF"

FONT = "Arial"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(hex_value: str) -> RGBColor:
    value = hex_value.strip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def set_run_font(run, size=None, color=None, bold=None, italic=None, all_caps=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if all_caps:
        rpr = run._element.get_or_add_rPr()
        caps = rpr.find(qn("w:caps"))
        if caps is None:
            caps = OxmlElement("w:caps")
            rpr.append(caps)
        caps.set(qn("w:val"), "1")


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=LIGHT_GRAY, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_paragraph_border_bottom(paragraph, color=PRIMARY_BLUE, size="12", space="4"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    old_grid = tbl.tblGrid
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    tbl.insert(0, grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def clear_paragraph(paragraph):
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def add_field(paragraph, instr: str, display: str = ""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = f" {instr} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr_el)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)
    return run


def insert_toc_field(paragraph):
    clear_paragraph(paragraph)
    add_field(paragraph, r'TOC \o "1-3" \h \z \u', "Table of contents will update automatically in Microsoft Word.")


def set_update_fields_on_open(docx_path: Path):
    tmp = Path(tempfile.mkdtemp(prefix="frs_fields_"))
    try:
        with zipfile.ZipFile(docx_path, "r") as zf:
            zf.extractall(tmp)
        settings = tmp / "word" / "settings.xml"
        parser = etree.XMLParser(remove_blank_text=False)
        tree = etree.parse(str(settings), parser)
        root = tree.getroot()
        update = root.find(f".//{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}updateFields")
        if update is None:
            update = etree.Element(f"{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}updateFields")
            root.insert(0, update)
        update.set(f"{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}val", "true")
        tree.write(str(settings), xml_declaration=True, encoding="UTF-8", standalone="yes")
        tmp_out = docx_path.with_suffix(".tmp.docx")
        with zipfile.ZipFile(tmp_out, "w", zipfile.ZIP_DEFLATED) as zf:
            for path in tmp.rglob("*"):
                if path.is_file():
                    zf.write(path, path.relative_to(tmp).as_posix())
        tmp_out.replace(docx_path)
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def crop_logo():
    image = Image.open(SOURCE_LOGO).convert("RGBA")
    region = image.crop((120, 940, 670, 1175))
    pixels = region.load()
    xs, ys = [], []
    for y in range(region.height):
        for x in range(region.width):
            r, g, b, a = pixels[x, y]
            if a and not (r > 242 and g > 242 and b > 242):
                xs.append(x)
                ys.append(y)
    if not xs:
        raise RuntimeError("Could not locate logo artwork in source image.")
    pad = 18
    box = (
        max(min(xs) - pad, 0),
        max(min(ys) - pad, 0),
        min(max(xs) + pad, region.width),
        min(max(ys) + pad, region.height),
    )
    cropped = region.crop(box)
    cropped.save(LOGO_CROP)


def setup_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(CHARCOAL)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.color.rgb = rgb(PRIMARY_BLUE if name != "Heading 3" else DEEP_NAVY)
        style.font.bold = True
        pf = style.paragraph_format
        pf.keep_with_next = True
        pf.line_spacing = 1.10

    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].paragraph_format.space_before = Pt(16)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)

    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)

    styles["Heading 3"].font.size = Pt(11.5)
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)

    for toc_style_name in ("TOC 1", "TOC 2", "TOC 3"):
        if toc_style_name in styles:
            style = styles[toc_style_name]
            style.font.name = FONT
            style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
            style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
            style.font.color.rgb = rgb(CHARCOAL)


def setup_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True


def remove_default_paragraph(container):
    if container.paragraphs:
        p = container.paragraphs[0]
        if not p.text:
            p._element.getparent().remove(p._element)


def add_header_footer(section):
    header = section.header
    remove_default_paragraph(header)
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [6500, 2860])
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, 0, 0, 0, 0)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
    left = table.cell(0, 0).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = left.add_run("Functional Requirements Specification")
    set_run_font(run, size=8.5, color=CHARCOAL, bold=True, all_caps=True)
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.add_run().add_picture(str(LOGO_CROP), width=Inches(1.35))
    rule = header.add_paragraph()
    rule.paragraph_format.space_before = Pt(1)
    rule.paragraph_format.space_after = Pt(0)
    set_paragraph_border_bottom(rule, LIGHT_GRAY, size="6", space="1")

    footer = section.footer
    remove_default_paragraph(footer)
    rule = footer.add_paragraph()
    rule.paragraph_format.space_after = Pt(3)
    set_paragraph_border_bottom(rule, LIGHT_GRAY, size="6", space="1")
    ft = footer.add_table(rows=1, cols=3, width=Inches(6.5))
    ft.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(ft, [3000, 3360, 3000])
    labels = [
        ("Synexis Digital", WD_ALIGN_PARAGRAPH.LEFT),
        ("Page ", WD_ALIGN_PARAGRAPH.CENTER),
        ("Version 1.0", WD_ALIGN_PARAGRAPH.RIGHT),
    ]
    for idx, (text, align) in enumerate(labels):
        cell = ft.cell(0, idx)
        set_cell_margins(cell, 0, 0, 0, 0)
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=8.5, color=CHARCOAL)
        if idx == 1:
            add_field(p, "PAGE", "1")
            r2 = p.add_run(" of ")
            set_run_font(r2, size=8.5, color=CHARCOAL)
            add_field(p, "NUMPAGES", "1")

    first_footer = section.first_page_footer
    remove_default_paragraph(first_footer)
    first_rule = first_footer.add_paragraph()
    first_rule.paragraph_format.space_after = Pt(3)
    set_paragraph_border_bottom(first_rule, LIGHT_GRAY, size="6", space="1")
    first = first_footer.add_table(rows=1, cols=3, width=Inches(6.5))
    first.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(first, [3000, 3360, 3000])
    first_labels = [
        ("Synexis Digital", WD_ALIGN_PARAGRAPH.LEFT),
        ("Page ", WD_ALIGN_PARAGRAPH.CENTER),
        ("Version 1.0", WD_ALIGN_PARAGRAPH.RIGHT),
    ]
    for idx, (text, align) in enumerate(first_labels):
        cell = first.cell(0, idx)
        set_cell_margins(cell, 0, 0, 0, 0)
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=8.5, color=CHARCOAL)
        if idx == 1:
            add_field(p, "PAGE", "1")
            r2 = p.add_run(" of ")
            set_run_font(r2, size=8.5, color=CHARCOAL)
            add_field(p, "NUMPAGES", "1")


def add_cover(doc, title, project_name, metadata):
    top = doc.add_paragraph()
    top.paragraph_format.space_after = Pt(10)
    top.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    top.add_run().add_picture(str(LOGO_CROP), width=Inches(2.35))

    accent = doc.add_table(rows=1, cols=2)
    accent.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(accent, [1550, 7810])
    for cell in accent.rows[0].cells:
        set_cell_margins(cell, 0, 0, 0, 0)
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(0)
    set_cell_shading(accent.cell(0, 0), PRIMARY_BLUE)
    set_cell_shading(accent.cell(0, 1), LIGHT_BLUE)
    for cell in accent.rows[0].cells:
        set_cell_border(cell, color=WHITE, size="0")

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(52)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(6)
    kr = kicker.add_run("TECHNICAL SPECIFICATION")
    set_run_font(kr, size=10.5, color=PRIMARY_BLUE, bold=True, all_caps=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(title)
    set_run_font(r, size=30, color=DARK_CHARCOAL, bold=True)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(24)
    sr = sub.add_run(project_name)
    set_run_font(sr, size=16, color=DEEP_NAVY, bold=True)

    meta_table = doc.add_table(rows=0, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(meta_table, [2600, 6760])
    rows = [
        ("Document Version:", metadata["Document Version"]),
        ("Status:", metadata["Status"]),
        ("Prepared For:", metadata["Prepared For"]),
        ("Document Type:", metadata["Document Type"]),
        ("Date:", "28 July 2026"),
        ("Prepared By:", "Synexis Digital"),
    ]
    for label, value in rows:
        cells = meta_table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    style_table(meta_table, label_value=True, header=False)

    tagline = doc.add_paragraph()
    tagline.paragraph_format.space_before = Pt(42)
    tagline.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tr = tagline.add_run("CONNECTING INNOVATION. DELIVERING IMPACT.")
    set_run_font(tr, size=8.5, color=PRIMARY_BLUE, bold=True, all_caps=True)
    doc.add_page_break()


def add_front_matter(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Table of Contents")
    set_run_font(r, size=16, color=PRIMARY_BLUE, bold=True)
    set_paragraph_border_bottom(p, PRIMARY_BLUE)
    toc = doc.add_paragraph()
    toc.paragraph_format.space_after = Pt(0)
    insert_toc_field(toc)
    doc.add_page_break()


def style_paragraph(paragraph, text):
    paragraph.paragraph_format.space_after = Pt(6)
    run_color = CHARCOAL
    run_bold = None
    if text == "Document Control":
        paragraph.style = "Heading 1"
        set_paragraph_border_bottom(paragraph, PRIMARY_BLUE)
        run_color = PRIMARY_BLUE
        run_bold = True
    elif text == "PART B — FUNCTIONAL REQUIREMENTS SPECIFICATION (FRS)":
        paragraph.style = "Heading 1"
        set_paragraph_border_bottom(paragraph, PRIMARY_BLUE)
        paragraph.paragraph_format.page_break_before = True
        run_color = PRIMARY_BLUE
        run_bold = True
    elif re.match(r"^B\.\d+\s", text):
        paragraph.style = "Heading 2"
        run_color = DEEP_NAVY
        run_bold = True
    elif text.startswith("Appendix —") or text == "Sign-off":
        paragraph.style = "Heading 1"
        set_paragraph_border_bottom(paragraph, PRIMARY_BLUE)
        paragraph.paragraph_format.page_break_before = True
        run_color = PRIMARY_BLUE
        run_bold = True
    else:
        paragraph.style = "Normal"
    for run in paragraph.runs:
        set_run_font(run, size=None, color=run_color, bold=run_bold)


def copy_paragraph(doc, src_paragraph):
    text = src_paragraph.text.strip()
    if not text:
        return None
    p = doc.add_paragraph()
    p.add_run(text)
    style_paragraph(p, text.strip())
    return p


def width_pattern_for_table(table):
    rows = len(table.rows)
    cols = len(table.columns)
    header = [cell.text.strip() for cell in table.rows[0].cells] if rows else []
    if cols == 3:
        if header == ["Role", "Name", "Date"]:
            return [2600, 4300, 2460]
        return [2400, 2500, 4460]
    if cols == 2:
        if rows and table.rows[0].cells[0].text.strip().endswith(":"):
            return [2600, 6760]
        if header == ["Field", "Detail"]:
            return [2500, 6860]
        if header == ["FR ID", "Requirement"]:
            return [1500, 7860]
        if header == ["Category", "Requirement"]:
            return [2100, 7260]
        if header == ["Term", "Definition"]:
            return [1800, 7560]
    return [CONTENT_WIDTH_DXA // cols] * cols


def fill_cell_from_text(cell, text, header=False, first_col=False):
    cell.text = ""
    paragraphs = text.split("\n") if text else [""]
    for idx, part in enumerate(paragraphs):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0 if header else 2)
        p.paragraph_format.line_spacing = 1.08
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if header or first_col else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(part)
        set_run_font(r, size=9.3 if not header else 9.5, color=WHITE if header else CHARCOAL, bold=header or first_col)


def style_table(table, label_value=False, header=True):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = width_pattern_for_table(table)
    set_table_geometry(table, widths)
    for ri, row in enumerate(table.rows):
        if ri == 0 and header:
            repeat_table_header(row)
        for ci, cell in enumerate(row.cells):
            is_header = header and ri == 0
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, top=100, bottom=100, start=130, end=130)
            set_cell_border(cell, LIGHT_GRAY, "6")
            if is_header:
                set_cell_shading(cell, PRIMARY_BLUE)
            elif label_value and ci == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            elif ri % 2 == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            else:
                set_cell_shading(cell, WHITE)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0 if is_header else 2)
                p.paragraph_format.line_spacing = 1.08
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_header or (ci == 0 and not label_value) else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_run_font(
                        run,
                        size=9.3 if not is_header else 9.5,
                        color=WHITE if is_header else CHARCOAL,
                        bold=is_header or (label_value and ci == 0),
                    )


def copy_table(doc, src_table):
    rows, cols = len(src_table.rows), len(src_table.columns)
    target = doc.add_table(rows=rows, cols=cols)
    for ri, row in enumerate(src_table.rows):
        for ci, cell in enumerate(row.cells):
            fill_cell_from_text(target.cell(ri, ci), cell.text, header=(ri == 0), first_col=(ci == 0))
    header = [cell.text.strip() for cell in src_table.rows[0].cells] if rows else []
    style_table(target, label_value=(header in (["Field", "Detail"], ["Term", "Definition"])), header=True)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return target


def iter_block_items(parent):
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    if hasattr(parent, "element"):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._tc
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def extract_metadata(paragraph_text):
    meta = {
        "Document Version": "1.0",
        "Status": "Draft for Review",
        "Prepared For": "CivicFlow Program (14-Phase Delivery Roadmap)",
        "Document Type": "FRS",
        "Original Line": paragraph_text,
    }
    patterns = {
        "Document Version": r"Document Version:\s*(.*?)\s+Status:",
        "Status": r"Status:\s*(.*?)\s+Prepared For:",
        "Prepared For": r"Prepared For:\s*(.*?)\s+Document Type:",
        "Document Type": r"Document Type:\s*(.*)$",
    }
    for key, pattern in patterns.items():
        match = re.search(pattern, paragraph_text)
        if match:
            meta[key] = match.group(1).strip()
    return meta


def add_body_from_source(doc, src_doc):
    # First three non-empty paragraphs are transformed into the designed cover.
    non_empty_seen = 0
    for block in iter_block_items(src_doc):
        if getattr(block, "text", None) is not None:
            text = block.text
            if text.strip():
                non_empty_seen += 1
            if non_empty_seen <= 3:
                continue
            if not text.strip():
                continue
            copy_paragraph(doc, block)
        else:
            copy_table(doc, block)


def apply_document_properties(doc, title, project):
    props = doc.core_properties
    props.title = title
    props.subject = project
    props.author = "Synexis Digital"
    props.comments = "Professionally redesigned FRS; original technical content preserved."


def build():
    crop_logo()
    src_doc = Document(str(SOURCE_DOCX))
    first_texts = [p.text.strip() for p in src_doc.paragraphs if p.text.strip()][:3]
    project_name, title, meta_line = first_texts
    metadata = extract_metadata(meta_line)

    doc = Document()
    setup_section(doc.sections[0])
    setup_styles(doc)
    add_header_footer(doc.sections[0])
    add_cover(doc, title, project_name, metadata)
    add_front_matter(doc)
    add_body_from_source(doc, src_doc)
    apply_document_properties(doc, title, project_name)
    doc.save(str(OUTPUT_DOCX))
    set_update_fields_on_open(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    build()
